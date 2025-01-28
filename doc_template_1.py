from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import parse_xml

# Create a new Document
doc = Document()

# Add a header and insert the logo
header = doc.sections[0].header
paragraph = header.paragraphs[0]

# Add the logo image to the header
logo_run = paragraph.add_run()
logo_run.add_picture("logo-dark.png", width=Inches(1))

# Remove the border line around the image
inline_shape = logo_run._element.xpath('.//a:blip')[0].getparent().getparent().getparent()
graphic_frame_properties = inline_shape.xpath('.//a:graphicFramePr')
if graphic_frame_properties:
    graphic_frame_properties[0].clear()  # Clear any frame properties, including borders

# Add header text
text_run = paragraph.add_run()
text_run.text = '\t' + "My Awesome Header"  # For center-align of text
text_run.style = "Heading 2 Char"
text_run.font.size = Pt(20)
text_run.bold = True
text_run.underline = True

# Adjust the spacing before the paragraph
paragraph_format = paragraph.paragraph_format
paragraph_format.space_before = Pt(0)

# Add temporary paragraphs for the meeting minutes content
doc.add_paragraph('Date: [Insert Date]')
doc.add_paragraph('Time: [Insert Time]')
doc.add_paragraph('Location: [Insert Location]')
doc.add_paragraph('Attendees: [Insert Attendees]')
doc.add_paragraph('Agenda: [Insert Agenda]')
doc.add_paragraph('Minutes: [Insert Minutes]')
doc.add_paragraph('Action Items: [Insert Action Items]')

# Save the document
doc.save('meeting_minutes_template.docx')
