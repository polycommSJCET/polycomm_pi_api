from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
import ast
import requests
import csv
import json
from datetime import datetime
import pytz
from docx2pdf import convert


# Variable to select the template
# TEMPLATE_SELECTION = 2  # Set to 1 for template 1, 2 for template 2
# m_id = 'sample_meeting'
SUPABASE_URL = "https://iljsvpxoiwnwxtjxypgi.supabase.co"
SUPABASE_KEY = "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6ImlsanN2cHhvaXdud3h0anh5cGdpIiwicm9sZSI6InNlcnZpY2Vfcm9sZSIsImlhdCI6MTczODk5NzYwMiwiZXhwIjoyMDU0NTczNjAyfQ.yHQOgRkEGo_NvJeT6ikONQQLiIhcuVVdsNcjZJE_9Rg"  # Use the service role key for secure upload
BUCKET_NAME = "Polycomm"  


def create_element(name):
    return OxmlElement(name)


def create_page_number(paragraph):
    run = paragraph.add_run()
    fldChar1 = create_element('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = create_element('w:instrText')
    instrText.text = "PAGE"
    fldChar2 = create_element('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def create_style(doc, name, font_name='Helvetica', size=11, color=RGBColor(0, 0, 0), bold=False):
    styles = doc.styles
    style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    font.name = font_name
    font.size = Pt(size)
    font.color.rgb = color
    font.bold = bold
    return style


def format_date(iso_str):
    utc_time = datetime.fromisoformat(iso_str)
    local_tz = pytz.timezone('Asia/Kolkata')
    local_time = utc_time.astimezone(local_tz)
    return local_time.strftime('%B %d, %Y')

def format_time(iso_str):
    utc_time = datetime.fromisoformat(iso_str)
    local_tz = pytz.timezone('Asia/Kolkata')
    local_time = utc_time.astimezone(local_tz)
    return local_time.strftime('%I:%M %p')


def fetch_meeting_info(meeting_id):
    # Fetch meeting info from the database
    response = requests.get(f"{SUPABASE_URL}/rest/v1/meetings?meeting_id=eq.{meeting_id}", headers={
        'apikey': SUPABASE_KEY,
        'Authorization': f'Bearer {SUPABASE_KEY}'
    })
    response.raise_for_status()
    meeting_info = response.json()
    if (meeting_info):
        return meeting_info[0]
    else:
        raise ValueError("Meeting info not found")


def create_document_template_1(m_id, organization_name, title, meeting_type, logo_path):
    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)

    styles = doc.styles
    style_h1 = styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
    font_h1 = style_h1.font
    font_h1.name = 'Calibri'
    font_h1.size = Pt(16)
    font_h1.bold = True
    font_h1.color.rgb = RGBColor(0, 51, 102)
    style_h1.paragraph_format.space_before = Pt(24)
    style_h1.paragraph_format.space_after = Pt(12)

    style_h2 = styles.add_style('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
    font_h2 = style_h2.font
    font_h2.name = 'Calibri'
    font_h2.size = Pt(14)
    font_h2.bold = True
    style_h2.paragraph_format.space_before = Pt(18)
    style_h2.paragraph_format.space_after = Pt(6)

    style_normal = styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Calibri'
    font_normal.size = Pt(11)
    style_normal.paragraph_format.space_after = Pt(12)
    style_normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = "CONFIDENTIAL"
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_para.style = styles['Normal']
    header_para.runs[0].font.bold = True
    header_para.runs[0].font.size = Pt(8)

    header_table = header.add_table(1, 2, width=Inches(6))
    header_cells = header_table.rows[0].cells
    logo_run = header_cells[0].paragraphs[0].add_run()

    company_para = header_cells[1].paragraphs[0]
    company_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    company_para.style = styles['Normal']
    if not company_para.runs:
        company_para.add_run()
    company_para.runs[0].text = organization_name
    company_para.runs[0].font.bold = True
    company_para.runs[0].font.size = Pt(14)

    doc_title = doc.add_heading('MEETING MINUTES', 0)
    doc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc_title.style = styles['Normal']
    doc_title.runs[0].font.size = Pt(18)
    doc_title.runs[0].font.bold = True

    meeting_info_table = doc.add_table(rows=6, cols=2)
    meeting_info_table.style = 'Table Grid'

    meeting_info = fetch_meeting_info(m_id)
    
    print(meeting_info)

    start_date = format_date(meeting_info.get('started_at', 'N/A'))
    start_time = format_time(meeting_info.get('started_at', 'N/A'))
    end_time = format_time(meeting_info.get('ended_at', 'N/A'))

    meeting_details = [
        ('Meeting Title:', title),
        ('Date:', start_date),
        ('Time:', f"{start_time} - {end_time}"),
        ('Meeting Called By:', meeting_info.get('ended_by_name', 'N/A')),
        ('Meeting Type:', meeting_type),
        ('Location:', 'Polycomm')
    ]

    for i, (label, value) in enumerate(meeting_details):
        cell_label = meeting_info_table.cell(i, 0)
        cell_value = meeting_info_table.cell(i, 1)
        cell_label.text = label
        cell_value.text = str(value)  # Ensure value is a string
        cell_label.paragraphs[0].runs[0].font.bold = True

    if logo_path:
        logo_run.add_picture(logo_path, width=Inches(1))

    doc.add_paragraph().add_run('\n')

    file_input=''

    with open('__temp__/csv/'+m_id+'.csv', mode='r') as file:
        csv_reader = csv.reader(file)
            
        for row in csv_reader:
            file_input += ", ".join(row) + "\n"

    response = requests.post(
        'http://localhost:11434/api/chat',
        json={
            "model": "llama3.2",
            "messages": [
                {
                    "role": "system",
                    "content": """Generate meeting minutes and reply in the following format and make sure the strings are closed on the same line, ensure the reply is parsable with ast.literal_eval(), reply only available datas:
       [('1. Meeting Overview', [
           ('sub heading 1', 'content 1'),        
           ('sub heading 2', 'content 2'),
        ]),
       ('2. Attendees', [
           ('Present', 'Names of present attendees'), 
           ('Apologies', 'Names of attendees who sent apologies'),
           ('Absent', 'Names of absent attendees')
       ]),
       ('3. Discussion Points', [
           ('Key Points Discussed',
            '1. Key Point 1\n' 
            '   - sub point 1\n' 
            '   - sub point 2\n'  
            '   - sub point n\n\n' 
            '2. Key Point 2\n' 
            '   - sub point 1\n' 
            '   - sub point 2\n\n' 
            '3. Key Point n\n' 
            '   - sub point 1\n' 
            '   - sub point n'),
           ('Decisions Made',
            '1. Decision 1\n' 
            '2. Decision 2\n' 
            '3. Decision n'),
           ('Voting Results',
            'result description')
       ]),
       ('4. Action Items', [
           ('Tasks Assigned',
            '1. task 1\n' 
            '2. task 2\n' 
            '3. task 3\n' 
            'n. task n'),
           ('Responsibilities',
            'Person 1: Responsibility\n' 
            'Person 2 : Responsibility\n' 
            'Person n: Responsibility'),
           ('Deadlines',
            'deadline date here')
       ]),
       ('5. Next Meeting', [
           ('Date and Time',
            'date and time mentioned'),
           ('Preliminary Agenda',
            '1. value 1\n' 
            '2. value 2\n' 
            '3. value 3\n' 
            '4. value 4')
       ])
    ]
                    """
                },
                {
                    "role": "user",
                    "content": "Generate meeting minutes based on the following conversation: "+file_input
                }
            ]
        },
        stream=True  # Enable streaming
    )

    message_content = []

    for line in response.iter_lines():

        if line:  # Ignore empty lines
            try:
                # Parse the line as JSON
                #print(line)
                data = json.loads(line)
                # Append content if the "message" key exists
                if "message" in data and "content" in data["message"]:
                    message_content.append(data["message"]["content"])
            except json.JSONDecodeError:
                print("Failed to decode line:", line)

    # Combine all message parts into the final response
    final_message = """""".join(message_content)

    start = final_message.find('[')
    end = final_message.rfind(']') + 1
    final_message = final_message[start:end]

    # Ensure the content is correctly parsed and added to the document
    try:
        sections_content = ast.literal_eval(final_message)
    except (SyntaxError, ValueError) as e:
        print(f"Error parsing final_message: {final_message}")
        raise e

    for section_title, subsections in sections_content:
        heading = doc.add_heading(section_title, level=1)
        heading.style = styles['CustomHeading1']
        for subsection in subsections:
            subheading = doc.add_heading(subsection[0], level=2)
            subheading.style = styles['CustomHeading2']
            content_para = doc.add_paragraph()
            content_para.style = styles['Normal']
            content_para.add_run(subsection[1])
            content_para.paragraph_format.left_indent = Inches(0.25)

    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.style = styles['Normal']
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    create_page_number(footer_para)
    footer_para.add_run(' | Confidential').font.size = Pt(8)

    doc.save('__temp__/docx/'+m_id+'_1.docx')
    convert('__temp__/docx/'+m_id+'_1.docx', '__temp__/pdf/'+m_id+'_1.pdf')


def create_document_template_2(m_id, organization_name, title, meeting_type, logo_path):
    print("template2")
    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    modern_title = create_style(doc, 'ModernTitle', 'Helvetica', 28, RGBColor(44, 62, 80), True)
    modern_subtitle = create_style(doc, 'ModernSubtitle', 'Helvetica', 14, RGBColor(52, 73, 94))
    modern_heading = create_style(doc, 'ModernHeading', 'Helvetica', 16, RGBColor(41, 128, 185), True)
    modern_subheading = create_style(doc, 'ModernSubheading', 'Helvetica', 12, RGBColor(52, 73, 94), True)
    modern_body = create_style(doc, 'ModernBody', 'Helvetica', 11, RGBColor(75, 75, 75))

    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = "CONFIDENTIAL"
    header_para.style = modern_subtitle
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    title_table = doc.add_table(rows=1, cols=2)
    title_table.autofit = False
    title_table.allow_autofit = False
    title_table.columns[0].width = Inches(2.0)
    title_table.columns[1].width = Inches(6.0)

    logo_cell = title_table.cell(0, 0)
    logo_run = logo_cell.paragraphs[0].add_run()

    title_cell = title_table.cell(0, 1)
    title_para = title_cell.paragraphs[0]
    title_para.style = modern_title
    title_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph('_' * 100).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph().add_run('\n')
    overview_para = doc.add_paragraph()
    overview_para.style = modern_heading
    overview_para.add_run('Meeting Overview').bold = True

    meeting_info = fetch_meeting_info(m_id)

    start_date = format_date(meeting_info.get('started_at', 'N/A'))
    start_time = format_time(meeting_info.get('started_at', 'N/A'))
    end_time = format_time(meeting_info.get('ended_at', 'N/A'))

    meeting_details = [
        ('Date & Time', f"{start_date} | {start_time} - {end_time}"),
        ('Location', 'Polycomm'),
        ('Organizer', meeting_info.get('ended_by_name', 'N/A')),
        ('Meeting ID', meeting_info.get('meeting_id', 'N/A'))
    ]

    for label, value in meeting_details:
        detail_para = doc.add_paragraph()
        detail_para.style = modern_body
        detail_para.paragraph_format.left_indent = Inches(0.5)
        detail_para.paragraph_format.space_after = Pt(6)
        detail_para.add_run(f'{label}: ').bold = True
        detail_para.add_run(value)

    if logo_path:
        logo_run.add_picture(logo_path, width=Inches(1.5))
    title_para.text = title

    format = """ [
        ('Participants', [
            ('Present', ['Person 1','Person 2']),
            ('Absent', ['Person 1','Person 2']),
        ]),
        ('Key Discussion Points', [
            ('Topic 1', [
                '• point 1',
                '• point 2',
                '• point 3'
            ]),
            ('Topic 2', [
                '• point 1',
                '• point 3',
                '• point n'
            ])
        ]),
        ('Action Items', [
            ('High Priority', [
                '→ Item 1 (Person assigned, Date)',
                '→ Item 2 (Person assigned, Date)',
                '→ Item 3 (Person assigned, Date)'
            ])
        ])
    ]"""

    file_input=''

    with open('__temp__/csv/'+m_id+'.csv', mode='r') as file:
        csv_reader = csv.reader(file)
            
        for row in csv_reader:
            file_input += ", ".join(row) + "\n"

    response = requests.post(
        'http://localhost:11434/api/chat',
        json={
            "model": "llama3.2",
            "messages": [
                {
                    "role": "system",
                    "content": """Generate meeting minutes and reply in the following format and make sure the strings are closed on the same line , ensure the reply is parsable with ast.literal_eval() , reply only available datas :
                    """+format
                },
                {
                    "role": "user",
                    "content": "Generate meeting minutes based on the following conversation: "+file_input
                }
            ]
        },
        stream=True  # Enable streaming
    )

    message_content = []

    for line in response.iter_lines():

        if line:  # Ignore empty lines
            try:
                # Parse the line as JSON
                #print(line)
                data = json.loads(line)
                # Append content if the "message" key exists
                if "message" in data and "content" in data["message"]:
                    message_content.append(data["message"]["content"])
            except json.JSONDecodeError:
                print("Failed to decode line:", line)

    # Combine all message parts into the final response
    final_message = """""".join(message_content)

    start = final_message.find('[')
    end = final_message.rfind(']') + 1
    final_message = final_message[start:end]

    # Ensure the content is correctly parsed and added to the document
    try:
        sections = ast.literal_eval(final_message)
    except (SyntaxError, ValueError) as e:
        print(f"Error parsing final_message: {final_message}")
        raise e

    for section_title, subsections in sections:
        section_para = doc.add_paragraph()
        section_para.style = modern_heading
        section_run = section_para.add_run(section_title)

        for subsection_title, items in subsections:
            subsection_para = doc.add_paragraph()
            subsection_para.style = modern_subheading
            subsection_para.paragraph_format.left_indent = Inches(0.3)
            subsection_para.paragraph_format.space_before = Pt(12)
            subsection_para.add_run(subsection_title)

            for item in items:
                item_para = doc.add_paragraph()
                item_para.style = modern_body
                item_para.paragraph_format.left_indent = Inches(0.6)
                item_para.paragraph_format.space_after = Pt(6)
                item_para.add_run(item)

        doc.add_paragraph().add_run('\n')
        separator = doc.add_paragraph()
        separator.add_run('・' * 50)
        separator.alignment = WD_ALIGN_PARAGRAPH.CENTER

    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.style = modern_subtitle
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    page_number = footer_para.add_run()
    create_page_number(footer_para)
    footer_para.add_run(f' | {organization_name}')

    doc.save('__temp__/docx/'+m_id+'_2.docx')
    convert('__temp__/docx/'+m_id+'_2.docx', '__temp__/pdf/'+m_id+'_2.pdf')


# if TEMPLATE_SELECTION == 1:
#       create_document_template_1("sample_meeting","Growisen","Weekly Meeting","Product development","logo-dark.png")
# elif TEMPLATE_SELECTION == 2:
#     create_document_template_2("sample_meeting")