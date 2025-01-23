import json
import requests
import csv
from fpdf import FPDF
from docx import Document
from datetime import datetime
import pandas as pd

class PDF(FPDF):
        def header(self):
            self.set_font("Arial", "B", 12)
            self.cell(0, 10, "Meeting Minutes", border=False, ln=True, align="C")
            self.ln(10)

        def footer(self):
            self.set_y(-15)
            self.set_font("Arial", "I", 8)
            self.cell(0, 10, f"Page {self.page_no()}", align="C")


def generate_pdf(content, m_id):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Arial", size=12)

    lines = content.split("\n")

    for line in lines:
        line = line.strip()

        # Handling Heading
        if line.startswith("<") and line.endswith(">"):  # Identifying headings (e.g., <Meeting Minutes>)
            pdf.set_font("Arial", "B", 14)
            pdf.ln(10)
            pdf.cell(0, 10, line.strip("<>"), ln=True)

        # Handling Bulleted Lists
        elif line.startswith("-") or line.startswith("*"):
            pdf.set_font("Arial", "I", 12)
            pdf.cell(10)  # Add indentation
            pdf.multi_cell(0, 10, line[2:].strip())  # Removing the bullet and adding a space

        # Handling Numbered Lists
        elif len(line) > 1 and line[0].isdigit() and line[1] == ".":
            pdf.set_font("Arial", "I", 12)
            pdf.cell(10)  # Add indentation
            pdf.multi_cell(0, 10, line.strip())

        # Handling Empty Lines
        elif line == "":
            pdf.ln(5)

        # Handling Regular Paragraphs
        else:
            pdf.set_font("Arial", size=12)
            pdf.multi_cell(0, 10, line)

    output_file = "__temp__/pdf/"+m_id + ".pdf"
    pdf.output(output_file)
    return output_file


def generate_word(content, m_id):
    # Create Word document
    doc = Document()
    doc.add_heading("Meeting Minutes", level=1)

    # Add generation date and time
    generation_date = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    doc.add_paragraph(f"Generated on: {generation_date}", style="Intense Quote")
    doc.add_paragraph("\n")  # Add spacing between header and body

    # Split and format content lines
    lines = content.split("\n")
    for line in lines:
        line = line.strip()

        # Handling Heading
        if line.startswith("<") and line.endswith(">"):  # Headings (e.g., <Heading>)
            clean_line = line.strip("<>")
            doc.add_heading(clean_line, level=2)

        # Handling Bulleted Lists
        elif line.startswith("-") or line.startswith("*"):
            clean_line = line[2:].strip()  # Removing bullet point and adding space
            doc.add_paragraph(clean_line, style="List Bullet")

        # Handling Numbered Lists
        elif len(line) > 1 and line[0].isdigit() and line[1] == ".":
            clean_line = line.strip()  # Removing leading digit
            doc.add_paragraph(clean_line, style="List Number")

        # Handling Empty Lines
        elif line == "":
            doc.add_paragraph("\n")  # Add empty space between sections

        # Handling Regular Paragraphs
        else:
            doc.add_paragraph(line)

    # Save the document
    output_file = f"__temp__/docx/{m_id}.docx"
    doc.save(output_file)
    return output_file

def generate_minutes(m_id):

    clear_duplicates(m_id)
    file_input=''

    with open('__temp__/csv/'+m_id+'.csv', mode='r') as file:
        csv_reader = csv.reader(file)
            
        for row in csv_reader:
            file_input += ", ".join(row) + "\n"



    response = requests.post(
        'http://localhost:11434/api/chat',
        json={
            "model": "llama3.2",
            "messages": [
                {
                    "role": "system",
                    "content": "Reply in format where each title/heading/sub-heading are enclosed in <angle brackets>."
                },
                {
                    "role": "user",
                    "content": "Generate meeting minutes based on the following conversation: "+file_input
                }
            ]
        },
        stream=True  # Enable streaming
    )

    # Collect the final message content
    message_content = []

    for line in response.iter_lines():
        if line:  # Ignore empty lines
            try:
                # Parse the line as JSON
                data = json.loads(line)
                # Append content if the "message" key exists
                if "message" in data and "content" in data["message"]:
                    message_content.append(data["message"]["content"])
            except json.JSONDecodeError:
                print("Failed to decode line:", line)

    # Combine all message parts into the final response
    final_message = "".join(message_content)
    #print(final_message)


    # Generate PDF
    output_file = generate_pdf(final_message,m_id)
    print(f"PDF generated and saved as {output_file}")


    return final_message


def clear_duplicates(m_id):

    file_path="__temp__/csv/"+m_id+".csv"
    df = pd.read_csv(file_path)

    print(len(df))

    i=0

    while(i<(len(df)-1)):
        #print(df.iloc[i]['speaker'])
        
        if((df.iloc[i]['speaker']==df.iloc[i+1]['speaker']) and (df.iloc[i]['translated_text'] in df.iloc[i+1]['translated_text'])):
            print("duplicate at index ",i)
            df.drop(i, axis=0, inplace=True)
            df.reset_index(drop=True, inplace=True)
            if i>0:
                i-=1


        else:
            i+=1

    df.to_csv(file_path, index=False)
    return file_path

